import os
import sys
import json
import datetime
import re
import matplotlib.pyplot as plt
from docx import Document
from pyannote.audio import Pipeline
import whisper
from langdetect import detect

# -----------------------
# Load models ONCE at startup
# -----------------------
hf_token = os.getenv("HUGGINGFACE_TOKEN")
if not hf_token:
    print("❌ HuggingFace token not found. Run: setx HUGGINGFACE_TOKEN your_token_here")
    sys.exit(1)

print("🔄 Loading Pyannote pipeline (first run may take time)...")
diarization_pipeline = Pipeline.from_pretrained(
    "pyannote/speaker-diarization",
    use_auth_token=hf_token
)

print("🔄 Loading Whisper model (first run may take time)...")
whisper_model = whisper.load_model("small")  # change to "tiny" or "base" for faster but less accurate

# -----------------------
# Utility: Create timestamped output folder
# -----------------------
def create_output_folder(base_name="run"):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    folder = os.path.join("outputs", f"{base_name}_{timestamp}")
    os.makedirs(folder, exist_ok=True)
    return folder

# -----------------------
# Save JSON
# -----------------------
def save_structured_json(data, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

# -----------------------
# Word transcript
# -----------------------
def generate_word_transcript(data, out_docx):
    doc = Document()
    doc.add_heading("Conversation Transcript", level=1)
    for seg in data["segments"]:
        speaker = seg.get("speaker", "Unknown")
        start = seg.get("startSec", 0)
        end = seg.get("endSec", 0)
        text = seg.get("text", "")
        doc.add_paragraph(f"{speaker} [{start:.2f}s - {end:.2f}s]: {text}")
    doc.save(out_docx)

# -----------------------
# Timeline plot
# -----------------------
def generate_timeline_image(data, out_png):
    plt.figure(figsize=(10, 2))
    y_labels = {}
    current_y = 0
    for seg in data["segments"]:
        speaker = seg["speaker"]
        if speaker not in y_labels:
            y_labels[speaker] = current_y
            current_y += 1
        plt.plot([seg["startSec"], seg["endSec"]],
                 [y_labels[speaker]]*2, linewidth=6)
    plt.yticks(list(y_labels.values()), list(y_labels.keys()))
    plt.xlabel("Time (s)")
    plt.tight_layout()
    plt.savefig(out_png)
    plt.close()

# -----------------------
# Code-switching Language Detection
# -----------------------
def detect_languages_in_text(text):
    if not text or not text.strip():
        return "unknown"
    chunks = re.split(r"[.?!]", text)
    detected = set()
    for chunk in chunks:
        chunk = chunk.strip()
        if len(chunk) < 2:
            continue
        try:
            lang_code = detect(chunk)
            detected.add(lang_code)
        except:
            detected.add("unknown")
    return ",".join(sorted(detected))

# -----------------------
# Improved Role Assignment
# -----------------------
def assign_roles(data):
    clinician_keywords = ["how long", "symptoms", "report", "examine", "check", "prescribe", "treatment"]
    patient_keywords = ["pain", "fever", "since", "problem", "hurt", "ache", "cough", "medicine"]

    for sp in data["speakers"]:
        spk_segments = [s for s in data["segments"] if s["speaker"] == sp["id"]]
        text_all = " ".join([s["text"] or "" for s in spk_segments]).lower()
        langs_all = ",".join([s["lang"] or "" for s in spk_segments])

        clin_score = sum(kw in text_all for kw in clinician_keywords)
        pat_score = sum(kw in text_all for kw in patient_keywords)

        lang_score_clin = langs_all.count("en")
        lang_score_pat = langs_all.count(",")  # code-switching indicator

        total_clin_score = clin_score + (0.5 * lang_score_clin)
        total_pat_score = pat_score + (0.5 * lang_score_pat)

        if total_clin_score > total_pat_score:
            sp["role"] = "clinician"
            sp["confidence"] = round(total_clin_score / max(1, len(text_all.split())), 2)
        elif total_pat_score > total_clin_score:
            sp["role"] = "patient"
            sp["confidence"] = round(total_pat_score / max(1, len(text_all.split())), 2)
        else:
            sp["role"] = "other"
            sp["confidence"] = 0.5

# -----------------------
# Main pipeline
# -----------------------
def main(audio_file=None, return_folder=False):
    # CLI fallback
    if audio_file is None and len(sys.argv) < 2:
        print("Usage: python final_pipeline.py <audio_file>")
        sys.exit(1)
    if audio_file is None:
        audio_file = sys.argv[1]

    output_folder = create_output_folder("diarization_run")

    # === 1. SPEAKER DIARIZATION ===
    diarization_result = diarization_pipeline(audio_file)

    speakers = []
    segments = []
    for turn, _, speaker in diarization_result.itertracks(yield_label=True):
        if speaker not in speakers:
            speakers.append(speaker)
        segments.append({
            "id": f"seg_{len(segments)+1:04d}",
            "speaker": speaker,
            "startSec": round(turn.start, 2),
            "endSec": round(turn.end, 2),
            "text": None,
            "lang": None
        })

    data = {
        "encounterId": "enc_001",
        "speakers": [{"id": s, "role": None, "confidence": None} for s in speakers],
        "segments": segments,
        "detectedLanguages": [],
        "createdAt": datetime.datetime.utcnow().isoformat() + "Z"
    }
    save_structured_json(data, os.path.join(output_folder, "output.json"))

    # === 2. SPEECH-TO-TEXT ===
    asr_result = whisper_model.transcribe(audio_file)
    for i, seg in enumerate(data["segments"]):
        if i < len(asr_result["segments"]):
            seg["text"] = asr_result["segments"][i]["text"].strip()

    # === 3. LANGUAGE DETECTION ===
    lang_set = set()
    for seg in data["segments"]:
        seg["lang"] = detect_languages_in_text(seg["text"])
        for l in seg["lang"].split(","):
            if l and l != "unknown":
                lang_set.add(l)
    data["detectedLanguages"] = sorted(list(lang_set))

    # === 4. ROLE ASSIGNMENT ===
    assign_roles(data)

    # === 5. SAVE OUTPUTS ===
    save_structured_json(data, os.path.join(output_folder, "final_output.json"))
    generate_word_transcript(data, out_docx=os.path.join(output_folder, "final_transcript.docx"))
    generate_timeline_image(data, out_png=os.path.join(output_folder, "timeline.png"))

    if return_folder:
        return output_folder

if __name__ == "__main__":
    main()
